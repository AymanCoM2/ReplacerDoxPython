from docx import Document
from docx2pdf import convert

document = Document('Invoice.docx')

dic = {
    '${a_1}': '1,380.00',
    '${a_2}': '23.00',
    '${a_3}': '0.00 %',
    '${a_4}': '23.00',
    '${a_5}': 'حبة E',
    '${a_6}': '60',
    '${a_7}': 'أيمن صلاح عبد الرازق يابياسيبلسقلصقصثقلث محمد صلاح علي مصطفى',
    # TODO , Replacing this For the Template
    '${a_8}': 'abcvvv',
    '${a_9}': '1',
}

for p in document.paragraphs:
    inline = p.runs
    for i in range(len(inline)):
        text = inline[i].text
        for key in dic.keys():
            if key in text:
                if key == '${a_7}':
                    # Split the text and insert line breaks
                    text = text.replace(key, dic[key])
                    lines = text.split('\n')
                    for j, line in enumerate(lines):
                        if j > 0:
                            p.add_run('\n')  # Add a line break
                            inline = p.runs  # Update the inline runs
                        inline[i + j].text = line
                else:
                    text = text.replace(key, dic[key])
                inline[i].text = text

document.save('new.docx')

convert("new.docx", "output.pdf")
